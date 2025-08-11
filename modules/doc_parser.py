import docx
import re

def extract_text(doc_file):
    try:
        doc = docx.Document(doc_file)
        return '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
    except Exception as e:
        return f"Error extracting text: {str(e)}"

def extract_clauses(doc_file):
    try:
        doc = docx.Document(doc_file)
        clauses = []
        current_clause = ""
        for para in doc.paragraphs:
            text = para.text.strip()
            if re.match(r"^\d+(\.\d+)*\s", text.lower()) or re.match(r"^clause\s+\d+", text.lower()):
                if current_clause:
                    clauses.append(current_clause.strip())
                current_clause = text
            else:
                current_clause += " " + text
        if current_clause:
            clauses.append(current_clause.strip())
        return clauses
    except Exception as e:
        return [f"Error extracting clauses: {str(e)}"]

def detect_doc_type(text):
    lower = text.lower()
    if "articles of association" in lower:
        return "Articles of Association"
    if "memorandum of association" in lower:
        return "Memorandum of Association"
    if "board resolution" in lower:
        return "Board Resolution"
    if "ubo declaration" in lower:
        return "UBO Declaration Form"
    if "register of members" in lower:
        return "Register of Members and Directors"
    if "license application" in lower:
        return "License Application Form"
    if "business plan" in lower:
        return "Business Plan"
    if "financial projections" in lower or "financial statements" in lower:
        return "Financial Statements"
    if "compliance manual" in lower:
        return "Compliance Manual"
    if "risk management" in lower:
        return "Risk Management Framework"
    if "annual compliance report" in lower:
        return "Annual Compliance Report"
    if "board meeting minutes" in lower:
        return "Board Meeting Minutes"
    if "risk assessment" in lower:
        return "Risk Assessment Report"
    if "compliance monitoring" in lower:
        return "Compliance Monitoring Report"
    return "Unknown"